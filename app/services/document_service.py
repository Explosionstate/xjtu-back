from __future__ import annotations

import shutil
from pathlib import Path
import re

from docx import Document as DocxDocument
from fastapi import UploadFile
from langchain_text_splitters import RecursiveCharacterTextSplitter
from PyPDF2 import PdfReader
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.core.config import settings
from app.core.errors import BusinessError
from app.models.document import Document, DocumentChunk
from app.models.knowledge_base import KnowledgeBase
from app.services.embedding_service import embed_texts
from app.services.embedding_service import normalize_embedding_model_name
from app.vectorstore.chroma_manager import delete_chunks_by_document, upsert_chunks

ALLOWED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


def _doc_storage_path(kb_id: str, doc_id: str, file_name: str) -> Path:
    return settings.docs_root / kb_id / f"{doc_id}_{file_name}"


def _extract_text_from_path(path: Path, ext: str) -> str:
    if ext in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if ext == ".pdf":
        reader = PdfReader(str(path))
        pages = [(page.extract_text() or "") for page in reader.pages]
        return "\n".join(pages)
    if ext == ".docx":
        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    raise BusinessError("不支持的文档格式", status_code=400)


def _split_text(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    cleaned_text = _clean_document_text(text)
    if not cleaned_text:
        return []
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""],
    )
    raw_chunks = [item.strip() for item in splitter.split_text(cleaned_text) if item.strip()]
    chunks: list[str] = []
    seen: set[str] = set()
    for item in raw_chunks:
        normalized = _normalize_chunk_text(item)
        if not normalized:
            continue
        key = normalized[:120].lower()
        if key in seen:
            continue
        seen.add(key)
        chunks.append(normalized)
    return chunks


def _clean_document_text(text: str) -> str:
    body = str(text or "").replace("\r\n", "\n")
    if not body.strip():
        return ""
    body = body.replace("\\.", ".").replace("\\-", "-").replace("\\+", "+")
    body = re.sub(r"(?m)^```.*$", "", body)
    body = re.sub(r"(?m)^#{1,6}\s*", "", body)
    body = re.sub(r"(?m)^\|?\s*[-:|]{3,}\s*\|?$", "", body)
    body = re.sub(r"(?m)^[-*•]\s*", "", body)
    body = re.sub(r"\n{3,}", "\n\n", body)
    return body.strip()


def _normalize_chunk_text(text: str) -> str:
    compact = str(text or "").replace("\r\n", "\n")
    compact = re.sub(r"(?m)^#{1,6}\s*", "", compact)
    compact = re.sub(r"(?m)^\|?\s*[-:|]{3,}\s*\|?$", "", compact)
    compact = re.sub(r"(?m)^[-*•]\s*", "", compact)
    compact = re.sub(r"(?m)^\d+[.)、]\s*", "", compact)
    compact = compact.replace("`", "")
    compact = re.sub(r"\s+", " ", compact).strip()
    compact = compact.strip(" ：:;；,，。-")
    if len(compact) < 12:
        return ""
    return compact


def _looks_like_mojibake(text: str) -> bool:
    body = (text or "").strip()
    if not body:
        return True
    replacement_ratio = body.count("�") / max(1, len(body))
    if replacement_ratio >= 0.03:
        return True
    markers = ["Ã", "Â", "ä¸", "ç", "å", "æ", "ï¼", "é"]
    marker_hits = sum(body.count(marker) for marker in markers)
    return marker_hits >= 16 and marker_hits / max(1, len(body)) > 0.06


def split_preview(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    return _split_text(text=text, chunk_size=chunk_size, chunk_overlap=chunk_overlap)


def upload_documents(
    db: Session,
    kb_id: str,
    files: list[UploadFile],
    chunk_size: int | None,
    chunk_overlap: int | None,
) -> list[Document]:
    kb = db.get(KnowledgeBase, kb_id)
    if kb is None or kb.status == "deleted":
        raise BusinessError("知识库不存在", status_code=404)

    target_chunk_size = chunk_size or settings.default_chunk_size
    target_chunk_overlap = chunk_overlap or settings.default_chunk_overlap
    saved_docs: list[Document] = []

    for upload in files:
        suffix = Path(upload.filename or "").suffix.lower()
        if suffix not in ALLOWED_EXTENSIONS:
            raise BusinessError(f"不支持的文件格式: {upload.filename}", status_code=400)

        doc = Document(
            kb_id=kb_id,
            file_name=upload.filename or "unknown",
            file_path="",
            file_type=suffix.lstrip("."),
            file_size=0,
            status="processing",
            chunk_count=0,
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)

        store_path = _doc_storage_path(kb_id, doc.id, upload.filename or "file")
        store_path.parent.mkdir(parents=True, exist_ok=True)
        try:
            with store_path.open("wb") as out:
                shutil.copyfileobj(upload.file, out)

            doc.file_path = store_path.as_posix()
            doc.file_size = store_path.stat().st_size

            text = _extract_text_from_path(store_path, suffix).strip()
            if not text or _looks_like_mojibake(text):
                doc.status = "failed"
                db.commit()
                db.refresh(doc)
                saved_docs.append(doc)
                continue

            chunks = _split_text(text, target_chunk_size, target_chunk_overlap)
            for idx, chunk in enumerate(chunks):
                db.add(
                    DocumentChunk(
                        kb_id=kb_id,
                        document_id=doc.id,
                        chunk_index=idx,
                        content=chunk,
                        source_location=doc.file_name,
                    )
                )
            db.commit()

            chunk_rows = list(
                db.scalars(
                    select(DocumentChunk)
                    .where(DocumentChunk.document_id == doc.id)
                    .order_by(DocumentChunk.chunk_index.asc())
                ).all()
            )
            upsert_chunks(
                kb_id=kb_id,
                chunk_ids=[item.id for item in chunk_rows],
                texts=[item.content for item in chunk_rows],
                metadatas=[
                    {
                        "kb_id": kb_id,
                        "document_id": doc.id,
                        "source_location": doc.file_name,
                        "chunk_index": item.chunk_index,
                    }
                    for item in chunk_rows
                ],
                embeddings=embed_texts(
                    [item.content for item in chunk_rows],
                    model_name=normalize_embedding_model_name(kb.embedding_model),
                ),
            )

            doc.chunk_count = len(chunk_rows)
            doc.status = "ready"
            db.commit()
            db.refresh(doc)
            saved_docs.append(doc)
        except Exception:
            doc.status = "failed"
            db.commit()
            db.refresh(doc)
            saved_docs.append(doc)

    return saved_docs


def list_documents(
    db: Session,
    kb_id: str,
    name: str | None,
    offset: int,
    limit: int,
) -> tuple[int, list[Document]]:
    stmt = select(Document).where(Document.kb_id == kb_id)
    count_stmt = (
        select(func.count()).select_from(Document).where(Document.kb_id == kb_id)
    )
    if name:
        stmt = stmt.where(Document.file_name.contains(name))
        count_stmt = count_stmt.where(Document.file_name.contains(name))
    stmt = stmt.order_by(Document.uploaded_at.desc()).offset(offset).limit(limit)
    total = db.scalar(count_stmt) or 0
    items = list(db.scalars(stmt).all())
    return total, items


def get_document_preview(
    db: Session, kb_id: str, document_id: str
) -> tuple[Document, str, list[str]]:
    doc = db.get(Document, document_id)
    if doc is None or doc.kb_id != kb_id:
        raise BusinessError("文档不存在", status_code=404)
    raw_text = (
        Path(doc.file_path).read_text(encoding="utf-8", errors="ignore")
        if doc.file_type in {"txt", "md"}
        else _extract_text_from_path(Path(doc.file_path), f".{doc.file_type}")
    )
    chunks = [
        chunk.content
        for chunk in db.scalars(
            select(DocumentChunk)
            .where(DocumentChunk.document_id == document_id)
            .order_by(DocumentChunk.chunk_index.asc())
        ).all()
    ]
    return doc, raw_text, chunks


def delete_document(db: Session, kb_id: str, document_id: str) -> None:
    doc = db.get(Document, document_id)
    if doc is None or doc.kb_id != kb_id:
        raise BusinessError("文档不存在", status_code=404)
    delete_chunks_by_document(kb_id=kb_id, document_id=document_id)
    db.query(DocumentChunk).filter(DocumentChunk.document_id == document_id).delete()
    if doc.file_path and Path(doc.file_path).exists():
        Path(doc.file_path).unlink(missing_ok=True)
    db.delete(doc)
    db.commit()


def delete_documents_batch(db: Session, kb_id: str, document_ids: list[str]) -> int:
    deleted = 0
    for document_id in document_ids:
        doc = db.get(Document, document_id)
        if doc is None or doc.kb_id != kb_id:
            continue
        delete_chunks_by_document(kb_id=kb_id, document_id=document_id)
        db.query(DocumentChunk).filter(
            DocumentChunk.document_id == document_id
        ).delete()
        if doc.file_path and Path(doc.file_path).exists():
            Path(doc.file_path).unlink(missing_ok=True)
        db.delete(doc)
        deleted += 1
    db.commit()
    return deleted


def reindex_document(
    db: Session,
    kb_id: str,
    document_id: str,
    chunk_size: int | None,
    chunk_overlap: int | None,
) -> Document:
    doc = db.get(Document, document_id)
    if doc is None or doc.kb_id != kb_id:
        raise BusinessError("文档不存在", status_code=404)
    kb = db.get(KnowledgeBase, kb_id)
    if kb is None or kb.status == "deleted":
        raise BusinessError("知识库不存在", status_code=404)
    text = _extract_text_from_path(Path(doc.file_path), f".{doc.file_type}")
    if _looks_like_mojibake(text):
        doc.status = "failed"
        doc.chunk_count = 0
        db.commit()
        db.refresh(doc)
        return doc
    target_chunk_size = chunk_size or settings.default_chunk_size
    target_chunk_overlap = chunk_overlap or settings.default_chunk_overlap

    delete_chunks_by_document(kb_id=kb_id, document_id=document_id)
    db.query(DocumentChunk).filter(DocumentChunk.document_id == document_id).delete()
    db.commit()

    chunks = _split_text(text, target_chunk_size, target_chunk_overlap)
    rows: list[DocumentChunk] = []
    for idx, chunk in enumerate(chunks):
        row = DocumentChunk(
            kb_id=kb_id,
            document_id=document_id,
            chunk_index=idx,
            content=chunk,
            source_location=doc.file_name,
        )
        db.add(row)
        rows.append(row)
    db.commit()

    rows = list(
        db.scalars(
            select(DocumentChunk)
            .where(DocumentChunk.document_id == document_id)
            .order_by(DocumentChunk.chunk_index.asc())
        ).all()
    )
    upsert_chunks(
        kb_id=kb_id,
        chunk_ids=[item.id for item in rows],
        texts=[item.content for item in rows],
        metadatas=[
            {
                "kb_id": kb_id,
                "document_id": document_id,
                "source_location": doc.file_name,
                "chunk_index": item.chunk_index,
            }
            for item in rows
        ],
        embeddings=embed_texts(
            [item.content for item in rows],
            model_name=normalize_embedding_model_name(kb.embedding_model),
        ),
    )
    doc.chunk_count = len(rows)
    doc.status = "ready"
    db.commit()
    db.refresh(doc)
    return doc
